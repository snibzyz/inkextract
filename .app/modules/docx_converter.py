import streamlit as st
from pathlib import Path
from typing import Dict, List
import os

from modules import paths


class DocxConverter:
    """แปลงไฟล์ DOCX เป็น TXT โดยรักษาโครงสร้าง subfolder"""

    def __init__(self):
        self.input_dir = paths.INPUT_DIR
        self.output_dir = paths.OUTPUT_DIR
        paths.ensure_dirs()
    
    def get_file_stats(self) -> Dict:
        """
        ดึงสถิติไฟล์ DOCX ในโฟลเดอร์ 0-input (รวม subfolder)
        
        Returns:
            Dict: สถิติไฟล์
        """
        stats = {
            'docx_files': 0,
            'txt_files': 0,
            'docx_size': 0,
            'txt_size': 0,
            'subfolders': 0
        }
        
        if not self.input_dir.exists():
            return stats
        
        # นับไฟล์ DOCX (รวม subfolder)
        docx_files = list(self.input_dir.rglob("*.docx"))
        stats['docx_files'] = len(docx_files)
        for f in docx_files:
            stats['docx_size'] += f.stat().st_size
        
        # นับไฟล์ TXT (รวม subfolder)
        txt_files = list(self.input_dir.rglob("*.txt"))
        stats['txt_files'] = len(txt_files)
        for f in txt_files:
            stats['txt_size'] += f.stat().st_size
        
        # นับ subfolder
        subfolders = [d for d in self.input_dir.rglob("*") if d.is_dir()]
        stats['subfolders'] = len(subfolders)
        
        return stats
    
    def get_docx_files_with_structure(self, source_dir: Path = None) -> List[Dict]:
        """
        ดึงรายการไฟล์ DOCX พร้อมโครงสร้าง subfolder
        
        Args:
            source_dir: โฟลเดอร์ต้นทาง (default: 0-input)
        
        Returns:
            List[Dict]: รายการไฟล์พร้อมข้อมูลโครงสร้าง
        """
        if source_dir is None:
            source_dir = self.input_dir
        
        if not source_dir.exists():
            return []
        
        files_info = []
        docx_files = list(source_dir.rglob("*.docx"))
        
        for file_path in docx_files:
            # คำนวณ relative path จาก source_dir
            relative_path = file_path.relative_to(source_dir)
            
            files_info.append({
                'full_path': file_path,
                'relative_path': relative_path,
                'parent_folder': relative_path.parent,
                'name': file_path.name,
                'size': file_path.stat().st_size
            })
        
        return files_info
    
    def convert_docx_to_txt(self, source_dir: Path = None, target_dir: Path = None, preserve_structure: bool = True) -> Dict:
        """
        แปลงไฟล์ DOCX เป็น TXT
        
        Args:
            source_dir: โฟลเดอร์ต้นทาง (default: 0-input)
            target_dir: โฟลเดอร์ปลายทาง (default: output)
            preserve_structure: รักษาโครงสร้าง subfolder หรือไม่
        
        Returns:
            Dict: ผลลัพธ์การแปลง
        """
        if source_dir is None:
            source_dir = self.input_dir
        if target_dir is None:
            target_dir = self.output_dir
        
        try:
            # ตรวจสอบว่ามี python-docx หรือไม่
            try:
                from docx import Document
            except ImportError:
                return {
                    'success': False,
                    'error': 'ไม่พบ python-docx library กรุณาติดตั้งด้วยคำสั่ง: pip install python-docx',
                    'files_processed': 0
                }
            
            # ตรวจสอบโฟลเดอร์ต้นทาง
            if not source_dir.exists():
                return {
                    'success': False,
                    'error': f"โฟลเดอร์ {source_dir} ไม่มีอยู่",
                    'files_processed': 0
                }
            
            # ค้นหาไฟล์ DOCX ทั้งหมด (รวม subfolder)
            docx_files_info = self.get_docx_files_with_structure(source_dir)
            
            if not docx_files_info:
                return {
                    'success': False,
                    'error': f"ไม่พบไฟล์ .docx ในโฟลเดอร์ {source_dir}",
                    'files_processed': 0
                }
            
            # แสดง progress bar
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            converted_files = []
            errors = []
            total_files = len(docx_files_info)
            
            for i, file_info in enumerate(docx_files_info):
                progress = (i + 1) / total_files
                progress_bar.progress(progress)
                
                file_path = file_info['full_path']
                relative_path = file_info['relative_path']
                
                status_text.text(f"กำลังแปลง: {relative_path} ({i+1}/{total_files})")
                
                try:
                    # อ่านไฟล์ DOCX
                    doc = Document(file_path)
                    
                    # แยกข้อความออกมา (เหมือนการ copy-paste)
                    full_text = []
                    
                    # อ่าน paragraph ทั้งหมด
                    for paragraph in doc.paragraphs:
                        text = paragraph.text.strip()
                        if text:  # เฉพาะบรรทัดที่มีเนื้อหา
                            full_text.append(text)
                    
                    # รวมข้อความ
                    content = '\n'.join(full_text)
                    
                    # กำหนดเส้นทางไฟล์ output
                    if preserve_structure:
                        # รักษาโครงสร้าง subfolder
                        output_file_path = target_dir / relative_path.parent / f"{file_path.stem}.txt"
                        # สร้าง subfolder ถ้ายังไม่มี
                        output_file_path.parent.mkdir(parents=True, exist_ok=True)
                    else:
                        # ไม่รักษาโครงสร้าง - เซฟที่ root ของ target_dir
                        output_file_path = target_dir / f"{file_path.stem}.txt"
                    
                    # บันทึกไฟล์ TXT
                    with open(output_file_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    
                    converted_files.append({
                        'source': str(relative_path),
                        'target': str(output_file_path.relative_to(target_dir)),
                        'paragraphs': len(doc.paragraphs),
                        'characters': len(content)
                    })
                    
                except Exception as e:
                    errors.append(f"❌ ไม่สามารถแปลง {relative_path}: {str(e)}")
            
            # เสร็จสิ้น
            progress_bar.progress(1.0)
            status_text.text("🎉 การแปลงเสร็จสิ้น!")
            
            return {
                'success': True,
                'files_processed': len(converted_files),
                'converted_files': converted_files,
                'errors': errors,
                'source_dir': str(source_dir),
                'target_dir': str(target_dir)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"เกิดข้อผิดพลาด: {str(e)}",
                'files_processed': 0
            }
    
    def preview_docx_files(self, source_dir: Path = None, max_files: int = 5) -> Dict:
        """
        แสดงตัวอย่างไฟล์ DOCX
        
        Args:
            source_dir: โฟลเดอร์ต้นทาง
            max_files: จำนวนไฟล์ที่แสดง
        
        Returns:
            Dict: ข้อมูลตัวอย่าง
        """
        if source_dir is None:
            source_dir = self.input_dir
        
        if not source_dir.exists():
            return {'error': f"โฟลเดอร์ {source_dir} ไม่มีอยู่"}
        
        try:
            from docx import Document
        except ImportError:
            return {'error': 'ไม่พบ python-docx library'}
        
        files_info = self.get_docx_files_with_structure(source_dir)
        
        if not files_info:
            return {'error': f"ไม่พบไฟล์ .docx ในโฟลเดอร์ {source_dir}"}
        
        previews = []
        
        for file_info in files_info[:max_files]:
            try:
                doc = Document(file_info['full_path'])
                
                # อ่าน 5 paragraph แรก
                preview_paragraphs = []
                for i, paragraph in enumerate(doc.paragraphs[:5]):
                    text = paragraph.text.strip()
                    if text:
                        preview_paragraphs.append(text)
                
                previews.append({
                    'name': str(file_info['relative_path']),
                    'size': file_info['size'],
                    'total_paragraphs': len(doc.paragraphs),
                    'preview': '\n'.join(preview_paragraphs),
                    'folder': str(file_info['parent_folder']) if file_info['parent_folder'] != Path('.') else 'root'
                })
            except Exception as e:
                previews.append({
                    'name': str(file_info['relative_path']),
                    'error': str(e)
                })
        
        return {
            'files': previews,
            'total_files': len(files_info),
            'directory': str(source_dir)
        }
    
    def convert_txt_to_docx(self, source_dir: Path = None, target_dir: Path = None) -> Dict:
        """
        แปลงไฟล์ TXT/MD เป็น DOCX

        Args:
            source_dir: โฟลเดอร์ต้นทาง (default: 2-clean)
            target_dir: โฟลเดอร์ปลายทาง (default: 2-clean — แปลง in-place)

        Returns:
            Dict: ผลลัพธ์การแปลง
        """
        if source_dir is None:
            source_dir = paths.CLEAN_DIR
        if target_dir is None:
            target_dir = paths.CLEAN_DIR  # in-place: เขียนใน 2-clean เดียวกัน
        
        # สร้างโฟลเดอร์ปลายทางถ้ายังไม่มี
        target_dir.mkdir(exist_ok=True)
        
        try:
            # ตรวจสอบว่ามี python-docx หรือไม่
            try:
                from docx import Document
            except ImportError:
                return {
                    'success': False,
                    'error': 'ไม่พบ python-docx library กรุณาติดตั้งด้วยคำสั่ง: pip install python-docx',
                    'files_processed': 0
                }
            
            # ตรวจสอบโฟลเดอร์ต้นทาง
            if not source_dir.exists():
                return {
                    'success': False,
                    'error': f"โฟลเดอร์ {source_dir} ไม่มีอยู่",
                    'files_processed': 0
                }
            
            # ค้นหาไฟล์ .txt และ .md ทั้งหมด
            txt_files = list(source_dir.glob("*.txt"))
            md_files = list(source_dir.glob("*.md"))
            all_files = txt_files + md_files
            
            if not all_files:
                return {
                    'success': False,
                    'error': f"ไม่พบไฟล์ .txt หรือ .md ในโฟลเดอร์ {source_dir}",
                    'files_processed': 0
                }
            
            # แสดง progress bar
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            converted_files = []
            errors = []
            total_files = len(all_files)
            
            for i, file_path in enumerate(all_files):
                progress = (i + 1) / total_files
                progress_bar.progress(progress)
                status_text.text(f"กำลังแปลง: {file_path.name} ({i+1}/{total_files})")
                
                try:
                    # อ่านไฟล์
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    # สร้างเอกสาร DOCX ใหม่
                    doc = Document()
                    
                    # แยกบรรทัด
                    lines = content.split('\n')
                    
                    # เพิ่มเนื้อหาแต่ละบรรทัด
                    for line in lines:
                        if line.strip():  # บรรทัดที่มีเนื้อหา
                            doc.add_paragraph(line)
                        else:  # บรรทัดว่าง
                            doc.add_paragraph('')
                    
                    # บันทึกไฟล์ DOCX
                    docx_file = target_dir / f"{file_path.stem}.docx"
                    doc.save(str(docx_file))
                    
                    converted_files.append({
                        'source': file_path.name,
                        'target': docx_file.name,
                        'lines': len(lines),
                        'characters': len(content)
                    })
                    
                except Exception as e:
                    errors.append(f"❌ ไม่สามารถแปลง {file_path.name}: {str(e)}")
            
            # เสร็จสิ้น
            progress_bar.progress(1.0)
            status_text.text("🎉 การแปลงเสร็จสิ้น!")
            
            return {
                'success': True,
                'files_processed': len(converted_files),
                'converted_files': converted_files,
                'errors': errors,
                'source_dir': str(source_dir),
                'target_dir': str(target_dir)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"เกิดข้อผิดพลาด: {str(e)}",
                'files_processed': 0
            }
    
    def get_txt_md_stats(self, source_dir: Path = None) -> Dict:
        """
        ดึงสถิติไฟล์ TXT/MD ในโฟลเดอร์ 2-clean
        
        Args:
            source_dir: โฟลเดอร์ต้นทาง (default: 2-clean)
        
        Returns:
            Dict: สถิติไฟล์
        """
        if source_dir is None:
            source_dir = paths.CLEAN_DIR

        stats = {
            'txt_files': 0,
            'md_files': 0,
            'total_files': 0,
            'total_size': 0
        }
        
        if not source_dir.exists():
            return stats
        
        # นับไฟล์ .txt
        txt_files = list(source_dir.glob("*.txt"))
        stats['txt_files'] = len(txt_files)
        for f in txt_files:
            stats['total_size'] += f.stat().st_size
        
        # นับไฟล์ .md
        md_files = list(source_dir.glob("*.md"))
        stats['md_files'] = len(md_files)
        for f in md_files:
            stats['total_size'] += f.stat().st_size
        
        stats['total_files'] = stats['txt_files'] + stats['md_files']
        
        return stats


